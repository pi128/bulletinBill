import os
import sqlite3
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
from slack_sdk import WebClient
from docx import Document

from dotenv import load_dotenv
load_dotenv()

# Load your Slack bot token and app-level token
SLACK_BOT_TOKEN = os.getenv("SLACK_BOT_TOKEN", "xoxb-your-token")
SLACK_APP_TOKEN = os.getenv("SLACK_APP_TOKEN", "xapp-your-app-token")

# Initialize Slack app
app = App(token=SLACK_BOT_TOKEN)
client = WebClient(token=SLACK_BOT_TOKEN)

# Database path
DB_PATH = "iac_bulletin.db"



@app.event("app_mention")
def handle_mention(event, say):
    print("Received mention event:", event)  # Debug log

    text = event["text"].lower()
    if "easter 3" in text:
        say("Here’s your bulletin for Easter 3 🙌")
        bulletin_path = generate_bulletin("Easter 3")
        upload_bulletin(event["channel"], bulletin_path)

    # Extract the liturgical day (naive method, improve later!)
    if "easter 3" in text:
        bulletin_path = generate_bulletin("Easter 3")
        upload_bulletin(event["channel"], bulletin_path)
        say("Here’s your bulletin for Easter 3 🙌")

def generate_bulletin(day_name):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute("SELECT id FROM calendar WHERE name = ?", (day_name,))
    result = cur.fetchone()
    if not result:
        return None
    cal_id = result[0]

    doc = Document()
    doc.add_heading(f"Bulletin for {day_name}", 0)

    # Collect
    cur.execute("SELECT content FROM collects WHERE calendar_id = ?", (cal_id,))
    collect = cur.fetchone()
    if collect:
        doc.add_heading("Collect", level=1)
        doc.add_paragraph(collect[0])

    # Readings
    cur.execute("SELECT type, reference FROM readings WHERE calendar_id = ?", (cal_id,))
    readings = cur.fetchall()
    if readings:
        doc.add_heading("Readings", level=1)
        for typ, ref in readings:
            doc.add_paragraph(f"{typ}: {ref}")

    # Save DOCX
    path = f"{day_name.replace(' ', '_').lower()}_bulletin.docx"
    doc.save(path)
    conn.close()
    return path

def upload_bulletin(channel_id, file_path):
    if not os.path.exists(file_path):
        return
    client.files_upload(
        channels=channel_id,
        file=file_path,
        title="Church Bulletin",
        filename=file_path
    )


if __name__ == "__main__":
    handler = SocketModeHandler(app, os.getenv("SLACK_APP_TOKEN"))
    handler.start()